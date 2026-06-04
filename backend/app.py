from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from werkzeug.utils import secure_filename
import re
import string
import os
import json
from collections import Counter
from io import StringIO, BytesIO
import numpy as np
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords

import torch
from transformers import AutoTokenizer, AutoModelForSequenceClassification

# File processing imports
import PyPDF2
import docx
import pandas as pd

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt', quiet=True)

try:
    nltk.data.find('tokenizers/punkt_tab')
except LookupError:
    nltk.download('punkt_tab', quiet=True)

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords', quiet=True)

app = Flask(__name__)
CORS(app)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Allowed extensions
ALLOWED_EXTENSIONS = {
    'pdf', 'docx', 'doc', 'pptx', 'ppt', 'txt', 'md', 'rtf', 'odt', 'tex',
    'xlsx', 'xls', 'csv', 'json',
    'png', 'jpg', 'jpeg', 'webp', 'gif', 'bmp',
    'py', 'js', 'html', 'css', 'java', 'cpp', 'cs', 'php', 'rb', 'ts', 'sh'
}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_file(file):
    """Extract text from various file formats."""
    filename = file.filename.lower()
    text = ""
    
    try:
        # Text files
        if filename.endswith('.txt') or filename.endswith('.md') or filename.endswith('.py') or \
           filename.endswith('.js') or filename.endswith('.html') or filename.endswith('.css') or \
           filename.endswith('.java') or filename.endswith('.cpp') or filename.endswith('.cs') or \
           filename.endswith('.php') or filename.endswith('.rb') or filename.endswith('.ts') or \
           filename.endswith('.sh') or filename.endswith('.json') or filename.endswith('.tex'):
            text = file.read().decode('utf-8', errors='ignore')
        
        # PDF files
        elif filename.endswith('.pdf'):
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        
        # Word documents
        elif filename.endswith('.docx') or filename.endswith('.doc'):
            doc = docx.Document(file)
            for para in doc.paragraphs:
                text += para.text + "\n"
        
        # CSV files
        elif filename.endswith('.csv'):
            df = pd.read_csv(file)
            text = df.to_string()
        
        # Excel files
        elif filename.endswith('.xlsx') or filename.endswith('.xls'):
            df = pd.read_excel(file)
            text = df.to_string()
        
        # JSON files
        elif filename.endswith('.json'):
            data = json.load(file)
            text = json.dumps(data, indent=2)
        
        # RTF files
        elif filename.endswith('.rtf'):
            text = file.read().decode('utf-8', errors='ignore')
            # Strip RTF formatting
            text = re.sub(r'\\[a-z]+\d*\s?', '', text)
            text = re.sub(r'[{}\\]', '', text)
        
        # ODT files
        elif filename.endswith('.odt'):
            doc = docx.Document(file)
            for para in doc.paragraphs:
                text += para.text + "\n"
        
        # PowerPoint (basic - limited support without additional libraries)
        elif filename.endswith('.pptx') or filename.endswith('.ppt'):
            text = "[PowerPoint files: Text extraction limited. Please copy-paste content.]"
        
        # Images - return message (OCR would require pytesseract installation)
        elif filename.endswith(('.png', '.jpg', '.jpeg', '.webp', '.gif', '.bmp')):
            text = "[Image files: Please copy-paste text content from images.]"
        
        else:
            text = "[Unsupported file format. Please try a text-based file.]"
    
    except Exception as e:
        text = f"[Error extracting text: {str(e)}]"
    
    return text

# Load the transformer-based detector
DETECTOR_MODEL_NAME = "Hello-SimpleAI/chatgpt-detector-roberta"
MAX_TOKENS = 512
CHUNK_STRIDE = 448  # 64-token overlap

print(f"Loading detector model: {DETECTOR_MODEL_NAME} ...")
tokenizer = AutoTokenizer.from_pretrained(DETECTOR_MODEL_NAME)
detector_model = AutoModelForSequenceClassification.from_pretrained(DETECTOR_MODEL_NAME)
detector_model.eval()
print("Detector model ready.")


def _score_chunk(input_ids, attention_mask):
    """Run one 512-token chunk through the model and return AI probability."""
    with torch.no_grad():
        outputs = detector_model(
            input_ids=input_ids.unsqueeze(0),
            attention_mask=attention_mask.unsqueeze(0),
        )
    probs = torch.softmax(outputs.logits, dim=-1)[0]
    # Label 1 = ChatGPT/AI in this model
    return float(probs[1].item())


def _predict_ai_probability(text):
    """Tokenize, chunk if needed, return mean AI probability in [0, 100]."""
    enc = tokenizer(text, return_tensors="pt", truncation=False)
    input_ids = enc["input_ids"][0]
    attention_mask = enc["attention_mask"][0]

    if input_ids.size(0) <= MAX_TOKENS:
        prob = _score_chunk(input_ids, attention_mask)
        return round(prob * 100, 2)

    chunk_probs = []
    for start in range(0, input_ids.size(0), CHUNK_STRIDE):
        end = min(start + MAX_TOKENS, input_ids.size(0))
        chunk_probs.append(_score_chunk(input_ids[start:end], attention_mask[start:end]))
        if end == input_ids.size(0):
            break
    mean_prob = sum(chunk_probs) / len(chunk_probs)
    return round(mean_prob * 100, 2)


def detect_ai_patterns(text):
    """Heuristic patterns shown to the user alongside the ML score."""
    patterns = []
    sentences = sent_tokenize(text)
    words = text.split()
    text_lower = text.lower()

    if len(sentences) > 3:
        sent_lens = [len(s.split()) for s in sentences]
        mean_len = sum(sent_lens) / len(sent_lens)
        if mean_len:
            variance = sum((x - mean_len) ** 2 for x in sent_lens) / len(sent_lens)
            cv = (variance ** 0.5) / mean_len
            if cv < 0.3:
                patterns.append("Unusually uniform sentence lengths")

    formal_phrases = [
        'furthermore', 'moreover', 'consequently', 'therefore', 'additionally',
        'in conclusion', 'it is important to note', 'it can be observed',
    ]
    if sum(1 for p in formal_phrases if p in text_lower) >= 2:
        patterns.append("Overly formal language structure")

    personal = {'i', 'my', 'me', 'we', 'our', 'us', 'you', 'your'}
    if words:
        pronoun_ratio = sum(1 for w in text_lower.split() if w in personal) / len(words)
        if pronoun_ratio < 0.02:
            patterns.append("Lack of personal language")

    return patterns


def analyze_text(text):
    """Main analysis function using the transformer detector."""
    text = text.strip()
    if len(text) < 50:
        return {
            'error': 'Text too short for analysis. Please provide at least 50 characters.',
            'ai_probability': None,
            'features': {},
            'patterns': []
        }

    try:
        ai_probability = _predict_ai_probability(text)
    except Exception as e:
        print(f"Detector error: {e}")
        return {'error': f'Detector failed: {e}'}

    patterns = detect_ai_patterns(text)

    if ai_probability >= 80:
        classification, confidence = "Likely AI-Generated", "High"
    elif ai_probability >= 60:
        classification, confidence = "Possibly AI-Generated", "Medium"
    elif ai_probability >= 40:
        classification, confidence = "Uncertain", "Low"
    elif ai_probability >= 20:
        classification, confidence = "Possibly Human-Written", "Medium"
    else:
        classification, confidence = "Likely Human-Written", "High"

    return {
        'ai_probability': ai_probability,
        'classification': classification,
        'confidence': confidence,
        'text_length': len(text),
        'word_count': len(text.split()),
        'sentence_count': len(sent_tokenize(text)),
        'features': {},
        'patterns': patterns,
        'ml_model_used': True,
        'model': 'chatgpt-detector-roberta',
    }

# Sample texts for demo
SAMPLE_TEXTS = {
    'ai': """The rapid advancement of artificial intelligence technologies has significantly transformed various sectors of modern society. This transformation encompasses multiple domains including healthcare, finance, education, and transportation. Furthermore, AI systems demonstrate remarkable capabilities in processing large datasets, identifying patterns, and generating predictions. Consequently, organizations increasingly adopt AI solutions to enhance operational efficiency and decision-making processes.""",
    
    'human': """I remember when I first started learning programming. It was challenging at first, but with practice and patience, I gradually improved. The key is to keep practicing and not be afraid to make mistakes. Every programmer writes buggy code initially; that's how we learn. What matters is persistence and the willingness to learn from errors. Looking back, I'm grateful for the struggles because they made me a better developer."""
}

@app.route('/analyze', methods=['POST'])
def analyze():
    """API endpoint to analyze text."""
    data = request.get_json()
    
    if not data or 'text' not in data:
        return jsonify({'error': 'No text provided'}), 400
    
    text = data['text']
    result = analyze_text(text)
    
    if 'error' in result:
        return jsonify(result), 400
    
    return jsonify(result)

@app.route('/upload', methods=['POST'])
def upload_file():
    """API endpoint to upload and extract text from files."""
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': f'File type not supported. Allowed: {", ".join(ALLOWED_EXTENSIONS)}'}), 400
    
    # Extract text from file
    text = extract_text_from_file(file)
    
    if not text or len(text.strip()) < 50:
        return jsonify({'error': 'Could not extract enough text from file. Please try a text-based file or copy-paste the content.'}), 400
    
    return jsonify({
        'filename': file.filename,
        'extracted_text': text[:10000],  # Limit to first 10000 chars
        'text_length': len(text)
    })

@app.route('/sample-text', methods=['GET'])
def get_sample_text():
    """API endpoint to get sample texts."""
    text_type = request.args.get('type', 'ai')
    
    if text_type not in SAMPLE_TEXTS:
        return jsonify({'error': 'Invalid sample type. Use "ai" or "human".'}), 400
    
    return jsonify({
        'type': text_type,
        'text': SAMPLE_TEXTS[text_type]
    })

def extract_claims(text):
    """Extract factual claims from text for verification."""
    # Split into sentences
    sentences = sent_tokenize(text)
    claims = []
    
    # Keywords that often indicate factual claims
    claim_indicators = [
        'according to', 'research shows', 'studies show', 'data indicates',
        'it is estimated', 'statistics show', 'experts say', 'research indicates',
        'the number of', 'percent of', 'increased by', 'decreased by',
        'since', 'because', 'therefore', 'as a result', 'consequently'
    ]
    
    for sentence in sentences:
        sentence_lower = sentence.lower()
        # Check if sentence contains claim indicators
        if any(indicator in sentence_lower for indicator in claim_indicators):
            # Check for specific factual patterns
            if any(pattern in sentence_lower for pattern in ['percent', '%', 'number', 'count', 'statistics', 'data']):
                claims.append({
                    'sentence': sentence.strip(),
                    'verified': None,
                    'sources': []
                })
            elif len(sentence.split()) > 10 and len(sentence.split()) < 50:
                # Substantial sentences may contain claims
                claims.append({
                    'sentence': sentence.strip(),
                    'verified': None,
                    'sources': []
                })
    
    # Return up to 5 claims
    return claims[:5]

def verify_claim(claim_text):
    """Verify a claim using web search (simplified mock for demo)."""
    # In production, use actual web search API
    # This is a simplified verification for demonstration
    
    import random
    import time
    
    # Simulate search delay
    time.sleep(0.3)
    
    # Common claim patterns to check (simplified)
    verified_claims = {
        'artificial intelligence': {
            'status': 'verified',
            'summary': 'AI is a rapidly growing field with applications across multiple industries.',
            'sources': ['Stanford AI Index Report 2024', 'MIT Technology Review']
        },
        'machine learning': {
            'status': 'verified',
            'summary': 'Machine learning is a subset of AI that enables systems to learn from data.',
            'sources': ['IEEE Computer Society', 'Nature Machine Intelligence']
        },
        'climate change': {
            'status': 'verified',
            'summary': 'Climate change is caused by greenhouse gas emissions from human activities.',
            'sources': ['IPCC Report', 'NASA Climate Research']
        },
        'healthcare': {
            'status': 'verified',
            'summary': 'AI in healthcare is transforming diagnosis, treatment, and patient care.',
            'sources': ['World Health Organization', 'NEJM AI']
        }
    }
    
    claim_lower = claim_text.lower()
    for key, value in verified_claims.items():
        if key in claim_lower:
            return value
    
    # Default response for unverified claims
    return {
        'status': 'unverified',
        'summary': 'Unable to verify this claim. Manual verification recommended.',
        'sources': []
    }

@app.route('/fact-check', methods=['POST'])
def fact_check():
    """API endpoint to fact-check claims in text."""
    data = request.get_json()
    
    if not data or 'text' not in data:
        return jsonify({'error': 'No text provided'}), 400
    
    text = data['text']
    
    # Extract claims from text
    claims = extract_claims(text)
    
    if not claims:
        return jsonify({
            'message': 'No specific claims detected for verification',
            'claims': []
        })
    
    # Verify each claim
    results = []
    for claim in claims:
        verification = verify_claim(claim['sentence'])
        results.append({
            'claim': claim['sentence'],
            'status': verification['status'],
            'summary': verification['summary'],
            'sources': verification['sources']
        })
    
    # Calculate verification stats
    verified_count = sum(1 for r in results if r['status'] == 'verified')
    unverified_count = sum(1 for r in results if r['status'] == 'unverified')
    
    return jsonify({
        'total_claims': len(results),
        'verified': verified_count,
        'unverified': unverified_count,
        'claims': results
    })

@app.route('/generate-report', methods=['POST'])
def generate_report():
    """API endpoint to generate PDF analysis report."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
    from reportlab.lib.units import inch
    import io
    import datetime
    
    data = request.get_json()
    
    if not data or 'analysis' not in data:
        return jsonify({'error': 'No analysis data provided'}), 400
    
    analysis = data['analysis']
    
    # Create PDF in memory
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
    
    # Styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=24, spaceAfter=30, textColor=colors.HexColor('#6366f1'))
    heading_style = ParagraphStyle('Heading', parent=styles['Heading2'], fontSize=14, spaceBefore=20, spaceAfter=10, textColor=colors.HexColor('#374151'))
    normal_style = ParagraphStyle('Normal', parent=styles['Normal'], fontSize=11, spaceAfter=10)
    
    # Build content
    story = []
    
    # Title
    story.append(Paragraph('AuthentiWrite', title_style))
    story.append(Paragraph('AI Detection Analysis Report', normal_style))
    story.append(Spacer(1, 20))
    
    # Date
    story.append(Paragraph(f"Report Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", normal_style))
    story.append(Spacer(1, 30))
    
    # Summary Section
    story.append(Paragraph('Analysis Summary', heading_style))
    
    # Create summary table
    summary_data = [
        ['Metric', 'Value'],
        ['AI Probability', f"{analysis.get('ai_probability', 0)}%"],
        ['Classification', analysis.get('classification', 'N/A')],
        ['Confidence', analysis.get('confidence', 'N/A')],
        ['Word Count', str(analysis.get('word_count', 0))],
        ['Sentence Count', str(analysis.get('sentence_count', 0))],
        ['Character Count', str(analysis.get('text_length', 0))]
    ]
    
    table = Table(summary_data, colWidths=[2.5*inch, 3*inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f3f4f6')),
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#6366f1')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e5e7eb'))
    ]))
    story.append(table)
    story.append(Spacer(1, 20))
    
    # Detection Results
    story.append(Paragraph('Detection Results', heading_style))
    
    # Determine result color
    prob = analysis.get('ai_probability', 0)
    if prob >= 70:
        result_text = "LIKELY AI-GENERATED"
        result_color = colors.HexColor('#ef4444')
    elif prob >= 50:
        result_text = "POSSIBLY AI-GENERATED"
        result_color = colors.HexColor('#f59e0b')
    else:
        result_text = "LIKELY HUMAN-WRITTEN"
        result_color = colors.HexColor('#22c55e')
    
    result_style = ParagraphStyle('Result', fontSize=16, textColor=result_color, spaceAfter=10)
    story.append(Paragraph(result_text, result_style))
    
    # Stylometric Features
    story.append(Paragraph('Stylometric Features', heading_style))
    features = analysis.get('features', {})
    if features:
        features_data = [
            ['Feature', 'Value'],
            ['Avg Sentence Length', f"{features.get('avg_sentence_length', 0):.2f}"],
            ['Avg Word Length', f"{features.get('avg_word_length', 0):.2f}"],
            ['Vocabulary Diversity', f"{features.get('vocabulary_diversity', 0):.2f}"],
            ['Punctuation Ratio', f"{features.get('punctuation_ratio', 0):.4f}"],
            ['Sentence Variation', f"{features.get('sentence_length_variation', 0):.4f}"],
            ['Long Word Ratio', f"{features.get('long_word_ratio', 0):.2f}"]
        ]
        
        f_table = Table(features_data, colWidths=[2.5*inch, 3*inch])
        f_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f9fafb')),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e5e7eb')),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e5e7eb')),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE')
        ]))
        story.append(f_table)
    
    # Detected Patterns
    patterns = analysis.get('patterns', [])
    if patterns:
        story.append(Paragraph('Detected Patterns', heading_style))
        for pattern in patterns:
            story.append(Paragraph(f"• {pattern}", normal_style))
    
    # Recommendations
    story.append(Paragraph('Recommendations', heading_style))
    
    if prob >= 70:
        story.append(Paragraph("This text shows strong indicators of AI generation. Further verification is recommended.", normal_style))
    elif prob >= 50:
        story.append(Paragraph("This text shows some AI-like characteristics. Review and consider requesting original drafts.", normal_style))
    else:
        story.append(Paragraph("This text appears to be human-written with low probability of AI generation.", normal_style))
    
    # Footer
    story.append(Spacer(1, 30))
    story.append(Paragraph("---", normal_style))
    story.append(Paragraph("AuthentiWrite - AI-Based Academic Integrity Detection System", normal_style))
    story.append(Paragraph("Federal University Oye-Ekiti | Department of Computer Science", normal_style))
    story.append(Paragraph("Project by Owofola Olakunle (FTP/CSC/25/0121610)", normal_style))
    
    # Build PDF
    doc.build(story)
    
    # Return PDF
    buffer.seek(0)
    return buffer.getvalue()

# Quality Analysis Functions
def calculate_readability(text):
    """Calculate readability scores using textstat."""
    import textstat
    scores = {
        'flesch_reading_ease': textstat.flesch_reading_ease(text),
        'flesch_kincaid_grade': textstat.flesch_kincaid_grade(text),
        'gunning_fog': textstat.gunning_fog(text),
        'smog_index': textstat.smog_index(text),
        'automated_readability_index': textstat.automated_readability_index(text),
        'coleman_liau_index': textstat.coleman_liau_index(text),
        'linsear_write_formula': textstat.linsear_write_formula(text),
        'dale_chall_readability_score': textstat.dale_chall_readability_score(text)
    }
    return scores

def detect_language(text):
    """Detect language of the text."""
    from langdetect import detect, LangDetectException
    try:
        lang = detect(text)
        lang_names = {
            'en': 'English', 'es': 'Spanish', 'fr': 'French', 'de': 'German',
            'it': 'Italian', 'pt': 'Portuguese', 'nl': 'Dutch', 'ru': 'Russian',
            'zh-cn': 'Chinese', 'ja': 'Japanese', 'ko': 'Korean', 'ar': 'Arabic',
            'hi': 'Hindi', 'tr': 'Turkish', 'pl': 'Polish', 'vi': 'Vietnamese',
            'th': 'Thai', 'sv': 'Swedish', 'da': 'Danish', 'fi': 'Finnish',
            'no': 'Norwegian', 'cs': 'Czech', 'el': 'Greek', 'he': 'Hebrew'
        }
        return {
            'language': lang,
            'language_name': lang_names.get(lang, lang.upper()),
            'confidence': 'high' if len(text) > 500 else 'medium' if len(text) > 200 else 'low'
        }
    except LangDetectException:
        return {'language': 'unknown', 'language_name': 'Unknown', 'confidence': 'low'}

def analyze_writing_style(text):
    """Analyze writing style for author attribution."""
    words = text.split()
    sentences = sent_tokenize(text)
    
    # Word-level features
    word_lengths = [len(w) for w in words]
    
    # Sentence-level features
    sentence_lengths = [len(s.split()) for s in sentences]
    
    # Vocabulary richness
    unique_words = len(set(words))
    total_words = len(words)
    
    # Punctuation patterns
    punctuation_marks = {
        'comma': text.count(','),
        'period': text.count('.'),
        'question': text.count('?'),
        'exclamation': text.count('!'),
        'semicolon': text.count(';'),
        'colon': text.count(':'),
        'quotes': text.count('"') + text.count("'"),
        'parentheses': text.count('(') + text.count(')'),
        'dash': text.count('-')
    }
    
    # Common words frequency
    stop_words = set(stopwords.words('english'))
    content_words = [w.lower() for w in words if w.lower() not in stop_words]
    content_word_freq = Counter(content_words)
    top_words = dict(content_word_freq.most_common(10))
    
    # Sentence starters (for style pattern)
    starters = [s.split()[0].lower() if s.split() else '' for s in sentences]
    starter_freq = Counter(starters)
    unique_starters = len(set(starters))
    
    return {
        'vocabulary_richness': round(unique_words / total_words if total_words else 0, 3),
        'avg_word_length': round(np.mean(word_lengths) if word_lengths else 0, 2),
        'avg_sentence_length': round(np.mean(sentence_lengths) if sentence_lengths else 0, 2),
        'sentence_length_variation': round(np.std(sentence_lengths) if sentence_lengths else 0, 2),
        'punctuation_usage': punctuation_marks,
        'top_content_words': top_words,
        'unique_sentence_starters': unique_starters,
        'writing_style_score': round((unique_words / total_words * 0.4 + unique_starters / len(sentences) * 0.6) if sentences and total_words else 0, 2)
    }

def check_plagiarism_similarity(text, database=None):
    """Check similarity with reference texts (simplified)."""
    # This is a simplified version - in production, you'd check against a database
    text_lower = text.lower()
    words = set(text_lower.split())
    
    # Common phrases that might indicate copying
    common_phrases = [
        'according to', 'as mentioned', 'as stated', 'it is important to note',
        'in conclusion', 'to summarize', 'as a result', 'due to the fact'
    ]
    
    phrase_matches = []
    for phrase in common_phrases:
        if phrase in text_lower:
            phrase_matches.append(phrase)
    
    # Check for very common academic phrases (low similarity risk)
    unique_ratio = len(words) / len(text.split()) if text.split() else 0
    
    # Simple similarity assessment
    similarity_risk = 'low'
    if len(phrase_matches) > 3:
        similarity_risk = 'medium'
    if unique_ratio < 0.3:
        similarity_risk = 'high'
    
    return {
        'similarity_risk': similarity_risk,
        'vocabulary_uniqueness': round(unique_ratio, 3),
        'common_phrase_matches': phrase_matches,
        'recommendation': 'Review content for originality' if similarity_risk != 'low' else 'Content appears original'
    }

@app.route('/quality-analysis', methods=['POST'])
def quality_analysis():
    """API endpoint for writing quality analysis."""
    data = request.get_json()
    
    if not data or 'text' not in data:
        return jsonify({'error': 'No text provided'}), 400
    
    text = data['text']
    
    if len(text) < 50:
        return jsonify({'error': 'Text too short for analysis (minimum 50 characters)'}), 400
    
    # Calculate all metrics
    readability = calculate_readability(text)
    language = detect_language(text)
    writing_style = analyze_writing_style(text)
    plagiarism = check_plagiarism_similarity(text)
    
    return jsonify({
        'readability': readability,
        'language': language,
        'writing_style': writing_style,
        'plagiarism_check': plagiarism,
        'text_length': len(text),
        'word_count': len(text.split())
    })

@app.route('/batch-analyze', methods=['POST'])
def batch_analyze():
    """API endpoint for batch processing multiple files."""
    if 'files' not in request.files:
        return jsonify({'error': 'No files provided'}), 400
    
    files = request.files.getlist('files')
    
    if len(files) == 0:
        return jsonify({'error': 'No files selected'}), 400
    
    if len(files) > 20:
        return jsonify({'error': 'Maximum 20 files allowed per batch'}), 400
    
    results = []
    
    for file in files:
        if file.filename == '':
            continue
        
        if not allowed_file(file.filename):
            results.append({
                'filename': file.filename,
                'error': 'Unsupported file type'
            })
            continue
        
        # Extract text
        text = extract_text_from_file(file)
        
        if len(text.strip()) < 50:
            results.append({
                'filename': file.filename,
                'error': 'Text too short for analysis'
            })
            continue
        
        # Analyze
        analysis = analyze_text(text)
        results.append({
            'filename': file.filename,
            'ai_probability': analysis['ai_probability'],
            'classification': analysis['classification'],
            'word_count': analysis['word_count']
        })
    
    # Summary statistics
    if results:
        probabilities = [r.get('ai_probability', 0) for r in results if 'ai_probability' in r]
        avg_probability = np.mean(probabilities) if probabilities else 0
        high_risk_count = sum(1 for p in probabilities if p >= 70)
    else:
        avg_probability = 0
        high_risk_count = 0
    
    return jsonify({
        'total_files': len(results),
        'analyzed_count': sum(1 for r in results if 'ai_probability' in r),
        'results': results,
        'summary': {
            'average_ai_probability': round(avg_probability, 2),
            'high_risk_count': high_risk_count,
            'low_risk_count': len(results) - high_risk_count
        }
    })

@app.route('/api/docs', methods=['GET'])
def api_docs():
    """API documentation for institutions."""
    return jsonify({
        'name': 'AuthentiWrite API',
        'version': '2.0.0',
        'description': 'AI-Based Academic Integrity Detection System API for Institutions',
        'authentication': {
            'type': 'API Key',
            'header': 'X-API-Key',
            'description': 'Include API key in request header for authenticated endpoints'
        },
        'endpoints': {
            'analyze': {
                'method': 'POST',
                'path': '/analyze',
                'description': 'Analyze text for AI generation probability',
                'body': {'text': 'string'},
                'response': {
                    'ai_probability': 'number (0-100)',
                    'classification': 'string',
                    'confidence': 'string',
                    'features': 'object',
                    'patterns': 'array'
                }
            },
            'upload': {
                'method': 'POST',
                'path': '/upload',
                'description': 'Upload file and extract text',
                'body': 'multipart/form-data with file',
                'response': {'extracted_text': 'string', 'filename': 'string'}
            },
            'quality_analysis': {
                'method': 'POST',
                'path': '/quality-analysis',
                'description': 'Analyze writing quality, readability, and style',
                'body': {'text': 'string'},
                'response': {
                    'readability': 'object with Flesch-Kincaid and other scores',
                    'language': 'object with detected language',
                    'writing_style': 'object with style features',
                    'plagiarism_check': 'object with similarity risk'
                }
            },
            'batch_analyze': {
                'method': 'POST',
                'path': '/batch-analyze',
                'description': 'Analyze multiple files at once',
                'body': 'multipart/form-data with files array (max 20)',
                'response': {
                    'results': 'array of per-file analysis',
                    'summary': 'aggregate statistics'
                }
            },
            'fact_check': {
                'method': 'POST',
                'path': '/fact-check',
                'description': 'Verify factual claims in text',
                'body': {'text': 'string'},
                'response': {'claims': 'array of verified claims'}
            },
            'generate_report': {
                'method': 'POST',
                'path': '/generate-report',
                'description': 'Generate PDF analysis report',
                'body': {'analysis': 'object from /analyze endpoint'},
                'response': 'PDF file download'
            }
        },
        'lms_integration': {
            'moodle': 'Use REST API endpoints with API key authentication',
            'canvas': 'Configure API endpoint URL and key in LMS settings',
            'blackboard': 'Integration via LTI or REST API'
        }
    })

@app.route('/health', methods=['GET'])
def health():
    """Health check endpoint."""
    return jsonify({
        'status': 'healthy',
        'model_loaded': MODEL_LOADED,
        'supported_formats': list(ALLOWED_EXTENSIONS)
    })

@app.route('/', methods=['GET'])
def index():
    """API info endpoint."""
    return jsonify({
        'name': 'AuthentiWrite API',
        'version': '1.2.0',
        'description': 'AI-Based Academic Integrity Detection System',
        'endpoints': {
            'analyze': 'POST /analyze - Analyze text for AI generation',
            'upload': 'POST /upload - Upload file for text extraction',
            'sample-text': 'GET /sample-text?type=ai|human - Get sample text',
            'fact-check': 'POST /fact-check - Verify claims in text',
            'generate-report': 'POST /generate-report - Generate PDF report',
            'health': 'GET /health - Check API health'
        },
        'supported_formats': {
            'documents': ['pdf', 'docx', 'doc', 'pptx', 'ppt', 'txt', 'md', 'rtf', 'odt', 'tex'],
            'spreadsheets': ['xlsx', 'xls', 'csv', 'json'],
            'images': ['png', 'jpg', 'jpeg', 'webp', 'gif', 'bmp'],
            'code': ['py', 'js', 'html', 'css', 'java', 'cpp', 'cs', 'php', 'rb', 'ts', 'sh']
        }
    })

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
# Test change at Thu Jun  4 23:35:04 UTC 2026
